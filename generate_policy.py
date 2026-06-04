"""
Generate a sample IC Reconciliation Policy DOCX for testing the policy upload feature.
Run: python generate_policy.py
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = r"C:\Users\kanne\Documents\DUMMY DATA"
os.makedirs(OUT_DIR, exist_ok=True)

BRAND_DARK  = RGBColor(0x1e, 0x1b, 0x4b)
BRAND_MID   = RGBColor(0x4f, 0x46, 0xe5)
GREY        = RGBColor(0x6b, 0x72, 0x80)
RED         = RGBColor(0xdc, 0x26, 0x26)
AMBER       = RGBColor(0xd9, 0x77, 0x06)
GREEN       = RGBColor(0x05, 0x96, 0x69)

doc = Document()

# ── page margins ─────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── style helpers ─────────────────────────────────────────────────────────────

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = BRAND_DARK
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = BRAND_MID
    return p

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BRAND_DARK
    return p

def body(text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        r = p.add_run(bold_prefix + "  ")
        r.bold = True
        r.font.color.rgb = BRAND_DARK
    p.add_run(text).font.size = Pt(10.5)
    return p

def bullet(text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after  = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix + " ")
        r.bold = True
        r.font.size = Pt(10.5)
    p.add_run(text).font.size = Pt(10.5)
    return p

def rule_box(label, text, color=BRAND_MID):
    """Inline rule highlight as a bold label + indented text."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r1 = p.add_run(f"RULE — {label}: ")
    r1.bold = True
    r1.font.color.rgb = color
    r1.font.size = Pt(10.5)
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    return p

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run("─" * 90)
    run.font.color.rgb = RGBColor(0xd1, 0xd5, 0xdb)
    run.font.size = Pt(8)


# ═══════════════════════════════════════════════════════════════════════════
# COVER
# ═══════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("NEXORA INTERNATIONAL GROUP")
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = BRAND_DARK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Intercompany Reconciliation Policy")
r.bold = True; r.font.size = Pt(16); r.font.color.rgb = BRAND_MID

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Version 3.2  |  Effective: 1 January 2026  |  Owner: Group Financial Controller")
r.font.size = Pt(10); r.font.color.rgb = GREY

doc.add_paragraph()
divider()
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════
# 1. PURPOSE AND SCOPE
# ═══════════════════════════════════════════════════════════════════════════
h1("1.  Purpose and Scope")
body(
    "This policy governs the end-to-end intercompany (IC) reconciliation process for all entities "
    "within the Nexora International Group. It sets out the standards, ownership rules, tolerance "
    "thresholds, SLA deadlines, and escalation criteria that must be followed at every period-end close."
)
body(
    "This policy applies to all entities that are consolidated into the Nexora Group financial statements, "
    "including wholly-owned subsidiaries, joint ventures with a Nexora majority interest, and any entity "
    "subject to an IC agreement with another Nexora entity."
)
rule_box("Scope", "All IC balances above USD 10,000 (or local currency equivalent) must be reconciled "
         "every period. Balances below this threshold are subject to annual reconciliation only.")

# ═══════════════════════════════════════════════════════════════════════════
# 2. ROLES AND RESPONSIBILITIES
# ═══════════════════════════════════════════════════════════════════════════
h1("2.  Roles and Responsibilities")

h2("2.1  Entity Controller")
bullet("Responsible for posting all IC entries in the entity's ERP within 3 business days of the transaction date.")
bullet("Responsible for confirming their entity's IC balances on the close calendar by the entity close deadline.")
bullet("First point of contact for all IC disputes where their entity is the 'buyer' or debtor in the IC agreement.")

h2("2.2  Group Financial Controller (GFC)")
bullet("Owns the reconciliation process and this policy.")
bullet("Reviews all open disputes exceeding USD 500,000 at day 3 of close.")
bullet("Has authority to waive tolerance thresholds for specific entity pairs with documented justification.")
bullet("Must approve any dispute carried forward to the following period.")

h2("2.3  Regional Finance Leads")
bullet("EMEA, Americas, and APAC Finance Leads are responsible for coordinating dispute resolution within their regions.")
bullet("Must escalate any dispute to the GFC if it cannot be resolved within the SLA window.")

rule_box("Ownership", "The 'seller' entity (the entity that raised the original IC charge or loan) "
         "is the default owner of any reconciliation dispute. Ownership may be transferred with written "
         "agreement from both entity controllers and approval from the relevant Regional Finance Lead.")

# ═══════════════════════════════════════════════════════════════════════════
# 3. TOLERANCE THRESHOLDS
# ═══════════════════════════════════════════════════════════════════════════
h1("3.  Tolerance Thresholds")
body(
    "A reconciliation pair is considered matched if the difference between the two entities' recorded "
    "balances falls within both of the following thresholds. Both conditions must be met simultaneously."
)

h2("3.1  Standard Tolerance (all entity pairs unless overridden)")
bullet("Absolute threshold:  USD 1,000 (or GBP 800 / EUR 920 equivalent)")
bullet("Percentage threshold:  0.1% of the larger of the two balances")
body("Differences within the standard tolerance are auto-matched and require no further action.")

h2("3.2  Elevated Tolerance (IC loans and long-term intercompany balances)")
bullet("Absolute threshold:  USD 5,000")
bullet("Percentage threshold:  0.05%")
body(
    "Elevated tolerance applies to accounts classified as IC Loan Receivable / Payable (account codes "
    "1600–1699 and 2600–2699). The lower percentage threshold reflects the materiality of these balances."
)

h2("3.3  Zero Tolerance (related-party transactions subject to transfer pricing documentation)")
bullet("Absolute threshold:  USD 0")
bullet("Any difference, however small, must be investigated and resolved before close sign-off.")
body(
    "Zero tolerance applies to all IC transactions that have been reviewed by the Group Tax team and are "
    "subject to an Advance Pricing Agreement (APA) or transfer pricing policy documentation."
)

rule_box("Threshold override",
         "The GFC may approve a one-time threshold increase for a specific entity pair by completing "
         "form GFC-IC-01. Approved overrides must be recorded in the IC Resolve tolerance configuration "
         "and are valid for one period only.", color=AMBER)

# ═══════════════════════════════════════════════════════════════════════════
# 4. FX TRANSLATION RULES
# ═══════════════════════════════════════════════════════════════════════════
h1("4.  FX Translation Rules")
body(
    "All IC balances must be translated to the Group reporting currency (USD) using the official "
    "period-end closing spot rates published by Group Treasury on the last business day of each month. "
    "No other rate source is permitted for period-end IC reconciliation purposes."
)

h2("4.1  Approved Rate Source")
bullet("Rate source:  Group Treasury / Bloomberg period-end closing rates")
bullet("Publication date:  Last business day of the period, by 17:00 London time")
bullet("Distribution:  Rates are published to the Group Finance SharePoint and loaded into IC Resolve")

h2("4.2  Rate Tolerance")
bullet("An FX rate difference of up to 0.5% versus the official rate is treated as a rounding difference and does not trigger an FX dispute.")
bullet("A rate difference exceeding 0.5% is classified as an FX Rate Error and must be corrected by the entity that used the incorrect rate.")

h2("4.3  Correction Procedure for FX Errors")
bullet("The entity that posted at the incorrect rate must repost using the official rate within 2 business days of the dispute being raised.")
bullet("If the correction cannot be posted within the SLA, the entity controller must notify the Regional Finance Lead and document the reason.")

rule_box("FX Rule", "Use of any rate other than the official Group Treasury period-end rate is a "
         "policy breach. Repeated FX rate errors (more than 2 per entity per year) must be reported "
         "to the Regional Finance Lead and documented in the entity's control environment.", color=RED)

# ═══════════════════════════════════════════════════════════════════════════
# 5. DISPUTE CLASSIFICATION AND SLA
# ═══════════════════════════════════════════════════════════════════════════
h1("5.  Dispute Classification and SLA Deadlines")
body(
    "All reconciliation disputes must be classified using the categories below. SLA deadlines are "
    "measured in business days from the date the dispute is raised in IC Resolve."
)

# table
tbl = doc.add_table(rows=6, cols=4)
tbl.style = "Table Grid"
hdrs = ["Dispute Type", "Definition", "Owner", "SLA (business days)"]
for i, h in enumerate(hdrs):
    cell = tbl.cell(0, i)
    cell.text = h
    run = cell.paragraphs[0].runs[0]
    run.bold = True
    run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
    cell._tc.get_or_add_tcPr().append(OxmlElement("w:shd"))
    shd = cell._tc.get_or_add_tcPr().find(qn("w:shd"))
    shd.set(qn("w:fill"), "1e1b4b")

data = [
    ("Missing Posting",    "One entity has no IC entry at all for the pair",
     "Entity with missing entry",  "1 day"),
    ("FX Rate Error",      "Difference caused by use of incorrect exchange rate",
     "Entity that used wrong rate", "2 days"),
    ("Timing Difference",  "Entry posted in a different period by one entity",
     "Entity that posted late",     "3 days"),
    ("Amount Difference",  "Underlying local currency amounts do not agree",
     "Seller entity",               "5 days"),
    ("Unknown / Mixed",    "Combination of the above or cannot be classified",
     "Group Financial Controller",  "3 days"),
]
for r, (dtype, defn, owner, sla) in enumerate(data, 1):
    tbl.cell(r, 0).text = dtype
    tbl.cell(r, 1).text = defn
    tbl.cell(r, 2).text = owner
    tbl.cell(r, 3).text = sla

doc.add_paragraph()

rule_box("SLA breach", "Any dispute that has not been resolved or formally escalated by its SLA deadline "
         "is automatically escalated to the Regional Finance Lead. Disputes unresolved after a further "
         "2 business days are escalated to the Group Financial Controller.", color=RED)

# ═══════════════════════════════════════════════════════════════════════════
# 6. ESCALATION CRITERIA
# ═══════════════════════════════════════════════════════════════════════════
h1("6.  Escalation Criteria")
body("The following conditions require immediate escalation to the Group Financial Controller, "
     "regardless of SLA status:")

bullet("Any single IC dispute exceeding USD 500,000")
bullet("Any entity pair where the combined unmatched exposure exceeds USD 1,000,000")
bullet("Any dispute involving a transfer pricing-documented transaction (zero tolerance)")
bullet("Any entity that has not confirmed its IC balances by its close deadline")
bullet("Any FX rate error that results in a misstatement exceeding USD 50,000")
bullet("Any dispute carried forward from the prior period that remains unresolved")

rule_box("Escalation",
         "Escalated disputes must be logged in IC Resolve with status 'Escalated'. "
         "The GFC must review all escalated disputes within 24 hours and either approve a resolution "
         "plan or authorise a prior-period adjustment.", color=RED)

# ═══════════════════════════════════════════════════════════════════════════
# 7. CLOSE CALENDAR AND CONFIRMATION
# ═══════════════════════════════════════════════════════════════════════════
h1("7.  Close Calendar and Entity Confirmation")

h2("7.1  Standard Close Timeline")
bullet("Day -3 before period end:  All IC entries must be posted in the ERP")
bullet("Day 0 (period end):  Trial balance extract uploaded to IC Resolve")
bullet("Day +2:  Matching run completed; disputes raised and assigned")
bullet("Day +4:  All disputes below USD 500k must be resolved or escalated")
bullet("Day +5:  Entity controllers confirm IC balances in IC Resolve")
bullet("Day +6:  GFC sign-off on reconciliation; close pack issued")

h2("7.2  Confirmation Requirements")
body(
    "Each entity controller must confirm their IC balances by marking their entity as 'Confirmed' "
    "in IC Resolve by Day +5. Confirmation means the entity controller attests that:"
)
bullet("All IC entries have been posted and are complete")
bullet("All disputes assigned to their entity have been resolved or formally escalated")
bullet("The trial balance extract matches the entity's ERP as at period end")

rule_box("Late confirmation",
         "An entity that has not confirmed by Day +5 will have its status flagged as 'Pending' "
         "and the Regional Finance Lead will be notified automatically. "
         "The GFC may proceed with group close without the entity's confirmation if the unmatched "
         "exposure is below the materiality threshold, subject to a documented exception.", color=AMBER)

# ═══════════════════════════════════════════════════════════════════════════
# 8. MATERIALITY AND CARRY-FORWARD
# ═══════════════════════════════════════════════════════════════════════════
h1("8.  Materiality and Carry-Forward")
body(
    "Disputes that cannot be resolved by the close deadline may be carried forward to the following "
    "period only if the following conditions are met:"
)
bullet("The dispute amount is below USD 250,000")
bullet("The dispute type is Timing Difference or FX Rate Error (not Amount Difference or Missing Posting)")
bullet("The entity controller has provided written confirmation of the expected resolution date")
bullet("The GFC has approved the carry-forward in IC Resolve")

body(
    "Any dispute exceeding USD 250,000, or any Missing Posting or Amount Difference dispute, "
    "must be resolved in the current period. If resolution is not possible, a prior-period adjustment "
    "or provision must be raised with approval from the CFO."
)

# ═══════════════════════════════════════════════════════════════════════════
# 9. AUDIT AND RECORD KEEPING
# ═══════════════════════════════════════════════════════════════════════════
h1("9.  Audit Trail and Record Keeping")
body(
    "IC Resolve maintains a complete audit trail of every match decision, dispute action, and AI-generated "
    "output. This audit trail must be preserved for a minimum of 7 years in accordance with the Group's "
    "document retention policy."
)
bullet("Every AI-generated dispute description is logged with the model version and timestamp")
bullet("Every tolerance override must reference the GFC-IC-01 approval form number")
bullet("Every escalation must be documented with the escalation reason and resolution outcome")
bullet("The close pack summary generated by IC Resolve constitutes part of the period-end close documentation")

# ═══════════════════════════════════════════════════════════════════════════
# 10. POLICY EXCEPTIONS AND AMENDMENTS
# ═══════════════════════════════════════════════════════════════════════════
h1("10.  Policy Exceptions and Amendments")
body(
    "Exceptions to this policy require written approval from the Group Financial Controller and must be "
    "documented in IC Resolve. Recurring exceptions for the same entity pair will trigger a policy review."
)
body(
    "This policy is reviewed annually by the Group Financial Controller. Amendments require sign-off "
    "from the CFO. All entities will be notified of material changes at least 30 days before they take effect."
)

divider()

# footer
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    "Nexora International Group  |  IC Reconciliation Policy v3.2  |  "
    "Approved by: CFO  |  Next review: January 2027  |  CONFIDENTIAL"
)
r.font.size = Pt(8.5)
r.font.color.rgb = GREY

# ── save ──────────────────────────────────────────────────────────────────────
path = os.path.join(OUT_DIR, "IC_Resolve_Policy_Sample.docx")
doc.save(path)
print(f"Saved: {path}")
