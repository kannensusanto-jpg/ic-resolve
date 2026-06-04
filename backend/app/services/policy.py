"""
Policy document management — extract text from PDF/DOCX/TXT uploads,
store in the DB, and retrieve the active policy for AI context injection.
"""
from io import BytesIO
from sqlalchemy.orm import Session
from app.models import PolicyDocument


def extract_text(file_bytes: bytes, filename: str) -> str:
    name = filename.lower()

    if name.endswith(".pdf"):
        import pypdf
        reader = pypdf.PdfReader(BytesIO(file_bytes))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n\n".join(p.strip() for p in pages if p.strip())

    if name.endswith(".docx"):
        import docx
        doc = docx.Document(BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)

    # Plain text / markdown
    for enc in ("utf-8", "latin-1"):
        try:
            return file_bytes.decode(enc)
        except UnicodeDecodeError:
            continue

    raise ValueError(f"Cannot extract text from {filename} — unsupported format")


def save_policy(db: Session, file_bytes: bytes, filename: str, label: str) -> PolicyDocument:
    text = extract_text(file_bytes, filename)
    if not text.strip():
        raise ValueError("Document appears to be empty or could not be parsed")

    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else "txt"

    doc = PolicyDocument(
        label=label,
        filename=filename,
        content_type=ext,
        raw_text=text,
        char_count=len(text),
        is_active=True,
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)
    return doc


def get_active_policy_text(db: Session) -> str | None:
    """Return the text of the most recently uploaded active policy, or None."""
    doc = (
        db.query(PolicyDocument)
        .filter(PolicyDocument.is_active == True)  # noqa: E712
        .order_by(PolicyDocument.uploaded_at.desc())
        .first()
    )
    return doc.raw_text if doc else None


def set_active(db: Session, policy_id: int, active: bool) -> PolicyDocument:
    doc = db.get(PolicyDocument, policy_id)
    if not doc:
        raise ValueError(f"Policy {policy_id} not found")
    doc.is_active = active
    db.commit()
    db.refresh(doc)
    return doc
